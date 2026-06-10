from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
REPORTS_DIR = ROOT / "docs" / "reports"
OUTPUT_ASCII = REPORTS_DIR / "requirements_report_fixed.docx"


def find_source_markdown() -> Path:
    matches = sorted(REPORTS_DIR.glob("*.md"))
    if not matches:
        raise FileNotFoundError("requirements markdown source not found")
    return matches[0]


def find_template_docx() -> Path:
    matches = [p for p in ROOT.rglob("*.docx") if "201" in p.name and "reports" not in {part.lower() for part in p.parts}]
    if not matches:
        raise FileNotFoundError("template docx not found")
    return sorted(matches)[0]


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_run_font(run, name: str = "SimSun") -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), name)


def set_cell_text(cell, text: str, align=None, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    set_run_font(run)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "auto")


def add_heading(doc: Document, text: str, level: int):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_paragraph(doc: Document, text: str, font_name: str = "SimSun"):
    paragraph = doc.add_paragraph(style="Normal")
    run = paragraph.add_run(text)
    set_run_font(run, font_name)
    return paragraph


def add_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = doc.tables[0].style
    set_table_borders(table)
    for index, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[index], value, WD_ALIGN_PARAGRAPH.CENTER, True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    return table


def parse_pipe_row(line: str) -> list[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [part.strip() for part in raw.split("|")]


def parse_cover_values(lines: list[str]) -> tuple[list[str], int]:
    values: list[str] = []
    if len(lines) < 3 or not lines[0].strip().startswith("|"):
        return values, 0
    index = 2
    while index < len(lines) and lines[index].strip().startswith("|"):
        row = parse_pipe_row(lines[index])
        if len(row) >= 2:
            values.append(row[1])
        index += 1
    return values, index


def fill_cover(doc: Document, values: list[str]) -> None:
    cover = doc.tables[0]
    for row_idx in range(min(len(values), len(cover.rows))):
        set_cell_text(cover.cell(row_idx, 1), values[row_idx], WD_ALIGN_PARAGRAPH.CENTER)


def clear_template_body(doc: Document) -> None:
    start_index = next(i for i, paragraph in enumerate(doc.paragraphs) if paragraph.style.name == "Heading 1")
    for paragraph in list(doc.paragraphs[start_index:]):
        delete_paragraph(paragraph)


def render_markdown(doc: Document, lines: list[str]) -> None:
    index = 0
    in_code = False
    skipped_title = False
    while index < len(lines):
        line = lines[index].rstrip("\n")
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            in_code = not in_code
            index += 1
            continue

        if in_code:
            add_paragraph(doc, line, "Consolas")
            index += 1
            continue

        if stripped.startswith("|"):
            block: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                block.append(parse_pipe_row(lines[index]))
                index += 1
            if len(block) >= 2:
                add_table(doc, [block[0]] + block[2:])
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            if level == 1 and not skipped_title:
                skipped_title = True
                index += 1
                continue
            add_heading(doc, text, min(level, 3))
            index += 1
            continue

        add_paragraph(doc, stripped)
        index += 1


def build() -> Path:
    markdown_path = find_source_markdown()
    template_path = find_template_docx()
    lines = markdown_path.read_text(encoding="utf-8-sig").splitlines()
    cover_values, body_start = parse_cover_values(lines)

    doc = Document(template_path)
    fill_cover(doc, cover_values)
    clear_template_body(doc)
    render_markdown(doc, lines[body_start:])
    doc.save(OUTPUT_ASCII)
    return OUTPUT_ASCII


if __name__ == "__main__":
    print(build())
