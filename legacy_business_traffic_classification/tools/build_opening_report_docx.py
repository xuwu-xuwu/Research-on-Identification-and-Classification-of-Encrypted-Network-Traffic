from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
SOURCE_MD = ROOT / "docs" / "reports" / "开题报告-网络加密数据流识别与分类技术研究.md"
TEMPLATE_DOCX = next((ROOT / "信息系统安全与对抗技术-报告模板").glob("*202*开题报告*.docx"))
OUTPUT_DOCX = ROOT / "docs" / "reports" / "开题报告-网络加密数据流识别与分类技术研究-模板规范版.docx"


BODY_CN_FONT = "宋体"
BODY_EN_FONT = "Times New Roman"
BODY_FONT_SIZE = Pt(12)
SUBHEAD_CN_FONT = "黑体"
SUBHEAD_EN_FONT = "Times New Roman"
SUBHEAD_FONT_SIZE = Pt(14)
TABLE_FONT_SIZE = Pt(10.5)
FIRST_LINE_INDENT = Cm(0.74)


def strip_inline_markdown(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("**", "")
    text = text.replace("*", "")
    text = text.replace("\u00a0", " ")
    return text.strip()


def starts_special(line: str) -> bool:
    stripped = line.strip()
    return bool(
        re.match(r"^(#{1,6})\s+", stripped)
        or stripped.startswith("|")
        or re.match(r"^\d+\.\s+", stripped)
        or re.match(r"^\[\d+\]", stripped)
    )


def parse_markdown(path: Path) -> list[dict]:
    lines = path.read_text(encoding="utf-8").splitlines()

    start = None
    for idx, line in enumerate(lines):
        if line.startswith("# 1 "):
            start = idx
            break
    if start is None:
        raise ValueError("Markdown source does not contain the first numbered section.")

    lines = lines[start:]
    blocks: list[dict] = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading_match.group(1)),
                    "text": strip_inline_markdown(heading_match.group(2)),
                }
            )
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            rows = []
            for row_line in table_lines:
                cells = [strip_inline_markdown(cell) for cell in row_line.strip().strip("|").split("|")]
                rows.append(cells)
            if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in rows[1]):
                header = rows[0]
                body = rows[2:]
            else:
                header = rows[0]
                body = rows[1:]
            blocks.append({"type": "table", "header": header, "rows": body})
            continue

        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                items.append(strip_inline_markdown(item))
                i += 1
            blocks.append({"type": "olist", "items": items})
            continue

        if re.match(r"^\[\d+\]", stripped):
            items = []
            while i < len(lines) and re.match(r"^\[\d+\]", lines[i].strip()):
                items.append(strip_inline_markdown(lines[i].strip()))
                i += 1
            blocks.append({"type": "refs", "items": items})
            continue

        para_lines = []
        while i < len(lines):
            current = lines[i].strip()
            if not current or starts_special(lines[i]):
                break
            para_lines.append(strip_inline_markdown(current))
            i += 1
        blocks.append({"type": "para", "text": " ".join(para_lines)})

    return blocks


def build_sections(blocks: list[dict]) -> dict[str, list[dict]]:
    sections: dict[str, list[dict]] = {}
    current_section = None

    for block in blocks:
        if block["type"] == "heading" and block["level"] == 1:
            heading_text = re.sub(r"^\d+\s+", "", block["text"]).strip()
            current_section = heading_text
            sections[current_section] = []
            continue

        if current_section is None:
            continue

        if block["type"] == "heading" and block["text"] == "结语":
            current_section = "时间计划"
            continue

        sections[current_section].append(block)

    return sections


def set_run_font(run, cn_font: str, en_font: str, size: Pt, bold: bool | None = None) -> None:
    run.font.name = en_font
    run.font.size = size
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), en_font)
    rfonts.set(qn("w:hAnsi"), en_font)
    rfonts.set(qn("w:eastAsia"), cn_font)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_paragraph_after(element, parent, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    paragraph = Paragraph(new_p, parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def insert_table_after(element, parent, document: Document, rows: int, cols: int) -> Table:
    table = document.add_table(rows=rows, cols=cols)
    tbl = table._tbl
    element.addnext(tbl)
    return Table(tbl, parent)


def format_body_paragraph(paragraph: Paragraph, first_line: bool = True) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = FIRST_LINE_INDENT if first_line else Cm(0)
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, BODY_CN_FONT, BODY_EN_FONT, BODY_FONT_SIZE)


def format_heading2_paragraph(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.25
    for run in paragraph.runs:
        set_run_font(run, SUBHEAD_CN_FONT, SUBHEAD_EN_FONT, SUBHEAD_FONT_SIZE, bold=True)


def format_list_paragraph(paragraph: Paragraph, index: int, text: str) -> None:
    paragraph.text = ""
    run = paragraph.add_run(f"{index}. {text}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.left_indent = FIRST_LINE_INDENT
    fmt.first_line_indent = -FIRST_LINE_INDENT
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    set_run_font(run, BODY_CN_FONT, BODY_EN_FONT, BODY_FONT_SIZE)


def format_reference_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.left_indent = FIRST_LINE_INDENT
    fmt.hanging_indent = FIRST_LINE_INDENT
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, BODY_CN_FONT, BODY_EN_FONT, BODY_FONT_SIZE)


def format_table(table: Table) -> None:
    for style_name in ("Table Grid", "TableGrid", "网格型", "网络型"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fmt = paragraph.paragraph_format
                fmt.first_line_indent = Cm(0)
                fmt.line_spacing = 1.2
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        BODY_CN_FONT if r_idx else SUBHEAD_CN_FONT,
                        BODY_EN_FONT,
                        TABLE_FONT_SIZE,
                        bold=(r_idx == 0),
                    )


def fill_cover(document: Document) -> None:
    if len(document.tables) < 2:
        return
    cover_table = document.tables[1]
    cover_table.cell(0, 1).text = "网络加密数据流识别与分类技术研究"

    for row_idx, row in enumerate(cover_table.rows):
        for col_idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                fmt = paragraph.paragraph_format
                fmt.first_line_indent = Cm(0)
                fmt.line_spacing = 1.25
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0)
                for run in paragraph.runs:
                    if col_idx == 0:
                        set_run_font(run, BODY_CN_FONT, BODY_EN_FONT, BODY_FONT_SIZE, bold=False)
                    else:
                        set_run_font(run, BODY_CN_FONT, BODY_EN_FONT, BODY_FONT_SIZE, bold=False)


def remove_placeholder_paragraphs(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        if idx <= 22:
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            continue
        if not paragraph.text.strip():
            delete_paragraph(paragraph)


def build_docx() -> None:
    blocks = parse_markdown(SOURCE_MD)
    sections = build_sections(blocks)

    document = Document(str(TEMPLATE_DOCX))
    fill_cover(document)
    remove_placeholder_paragraphs(document)

    heading_map: dict[str, Paragraph] = {}
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style and paragraph.style.name == "Heading 1" and text:
            heading_map[text] = paragraph

    for section_title, section_blocks in sections.items():
        if section_title not in heading_map:
            continue

        heading_para = heading_map[section_title]
        parent = heading_para._parent
        current_element = heading_para._p

        for block in section_blocks:
            if block["type"] == "heading" and block["level"] == 2:
                sub_para = insert_paragraph_after(current_element, parent, block["text"])
                format_heading2_paragraph(sub_para)
                current_element = sub_para._p
                continue

            if block["type"] == "para":
                para = insert_paragraph_after(current_element, parent, block["text"], style="Normal")
                format_body_paragraph(para, first_line=True)
                current_element = para._p
                continue

            if block["type"] == "olist":
                for idx, item in enumerate(block["items"], start=1):
                    para = insert_paragraph_after(current_element, parent, style="Normal")
                    format_list_paragraph(para, idx, item)
                    current_element = para._p
                continue

            if block["type"] == "refs":
                for item in block["items"]:
                    para = insert_paragraph_after(current_element, parent, item, style="Normal")
                    format_reference_paragraph(para)
                    current_element = para._p
                continue

            if block["type"] == "table":
                rows = [block["header"], *block["rows"]]
                col_count = max(len(row) for row in rows)
                table = insert_table_after(current_element, parent, document, rows=len(rows), cols=col_count)
                for r_idx, row in enumerate(rows):
                    for c_idx, value in enumerate(row):
                        table.cell(r_idx, c_idx).text = value
                format_table(table)
                current_element = table._tbl
                continue

    document.save(str(OUTPUT_DOCX))


if __name__ == "__main__":
    build_docx()
