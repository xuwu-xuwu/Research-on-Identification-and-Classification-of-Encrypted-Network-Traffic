from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORTS_DIR = ROOT / "docs" / "reports"
SOURCE_MD = REPORTS_DIR / "project_development_report_source.md"
TEMPLATE_DOCX = next((ROOT / "docs" / "templates").rglob("*401*.docx"))
OUTPUT_DOCX = REPORTS_DIR / "项目研制报告-网络流量加密方法识别与分类技术研究-模板规范版.docx"

BODY_CN_FONT = "宋体"
BODY_EN_FONT = "Times New Roman"
BODY_FONT_SIZE = Pt(12)
SUBHEAD_CN_FONT = "黑体"
SUBHEAD_EN_FONT = "Times New Roman"
SUBHEAD_FONT_SIZE = Pt(14)
TABLE_FONT_SIZE = Pt(9.5)
FIRST_LINE_INDENT = Cm(0.74)


def strip_inline_markdown(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("**", "")
    text = text.replace("*", "")
    text = text.replace("\u00a0", " ")
    return text.strip()


def parse_pipe_row(line: str) -> list[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [strip_inline_markdown(part) for part in raw.split("|")]


def parse_cover_values(lines: list[str]) -> tuple[list[str], int]:
    values: list[str] = []
    if len(lines) < 3 or not lines[0].strip().startswith("|"):
        return values, 0
    index = 2
    while index < len(lines) and lines[index].strip().startswith("|"):
        row = parse_pipe_row(lines[index])
        if len(row) >= 2:
            values.append(row[1])
        index += 1
    return values, index


def starts_special(line: str) -> bool:
    stripped = line.strip()
    return bool(
        re.match(r"^(#{1,6})\s+", stripped)
        or stripped.startswith("|")
        or re.match(r"^\d+\.\s+", stripped)
    )


def parse_markdown(lines: list[str]) -> list[dict]:
    blocks: list[dict] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading_match.group(1)),
                    "text": strip_inline_markdown(heading_match.group(2)),
                }
            )
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            rows = [parse_pipe_row(row_line) for row_line in table_lines]
            if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in rows[1]):
                header = rows[0]
                body = rows[2:]
            else:
                header = rows[0]
                body = rows[1:]
            blocks.append({"type": "table", "header": header, "rows": body})
            continue

        if re.match(r"^\d+\.\s+", stripped):
            items: list[str] = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                items.append(strip_inline_markdown(item))
                i += 1
            blocks.append({"type": "olist", "items": items})
            continue

        para_lines: list[str] = []
        while i < len(lines):
            current = lines[i].strip()
            if not current or starts_special(lines[i]):
                break
            para_lines.append(strip_inline_markdown(current))
            i += 1
        blocks.append({"type": "para", "text": " ".join(para_lines)})

    return blocks


def set_run_font(run, cn_font: str, en_font: str, size: Pt, bold: bool | None = None) -> None:
    run.font.name = en_font
    run.font.size = size
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key, value in (
        ("w:ascii", en_font),
        ("w:hAnsi", en_font),
        ("w:eastAsia", cn_font),
        ("w:cs", en_font),
    ):
        rfonts.set(qn(key), value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, align=WD_ALIGN_PARAGRAPH.CENTER, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, BODY_CN_FONT, BODY_EN_FONT, TABLE_FONT_SIZE, bold=bold)
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.left_indent = Cm(0)
    fmt.line_spacing = 1.15
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def fill_cover(doc: Document, values: list[str]) -> None:
    info_table = doc.tables[0]
    for row_idx, value in enumerate(values):
        if row_idx < len(info_table.rows):
            set_cell_text(info_table.cell(row_idx, 1), value)


def paragraph_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t"))).strip()


def clear_template_body(doc: Document) -> None:
    body = doc.element.body
    children = list(body)
    start_index = None
    for index, child in enumerate(children):
        if child.tag == qn("w:p") and paragraph_text(child) == "文档介绍":
            start_index = index
            break
    if start_index is None:
        raise ValueError("Could not locate the start of the template body.")

    for child in children[start_index:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def format_body_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = FIRST_LINE_INDENT
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, BODY_CN_FONT, BODY_EN_FONT, BODY_FONT_SIZE)


def format_heading_paragraph(paragraph, level: int) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.left_indent = Cm(0)
    fmt.space_before = Pt(12 if level == 1 else 6)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.25
    font_size = Pt(15 if level == 1 else 14)
    for run in paragraph.runs:
        set_run_font(run, SUBHEAD_CN_FONT, SUBHEAD_EN_FONT, font_size, bold=True)


def format_list_paragraph(paragraph, index: int, text: str) -> None:
    paragraph.text = ""
    run = paragraph.add_run(f"{index}. {text}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.left_indent = FIRST_LINE_INDENT
    fmt.first_line_indent = -FIRST_LINE_INDENT
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    set_run_font(run, BODY_CN_FONT, BODY_EN_FONT, BODY_FONT_SIZE)


def apply_table_widths(table, column_count: int) -> None:
    width_maps = {
        2: [4.2, 11.4],
        3: [4.0, 5.0, 6.6],
        4: [3.2, 4.0, 3.0, 5.4],
        5: [3.0, 3.2, 2.4, 2.4, 4.6],
    }
    widths = width_maps.get(column_count, [15.6 / column_count] * column_count)
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Cm(width)


def format_table(table) -> None:
    for style_name in ("Table Grid", "TableGrid", "网格型"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    apply_table_widths(table, len(table.columns))
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fmt = paragraph.paragraph_format
                fmt.first_line_indent = Cm(0)
                fmt.left_indent = Cm(0)
                fmt.line_spacing = 1.15
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        BODY_CN_FONT,
                        BODY_EN_FONT,
                        TABLE_FONT_SIZE,
                        bold=(r_idx == 0),
                    )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def render_markdown(doc: Document, blocks: list[dict]) -> None:
    for block in blocks:
        if block["type"] == "heading":
            paragraph = doc.add_paragraph()
            paragraph.add_run(block["text"])
            format_heading_paragraph(paragraph, level=block["level"])
            continue

        if block["type"] == "para":
            paragraph = doc.add_paragraph()
            paragraph.add_run(block["text"])
            format_body_paragraph(paragraph)
            continue

        if block["type"] == "olist":
            for index, item in enumerate(block["items"], start=1):
                paragraph = doc.add_paragraph()
                format_list_paragraph(paragraph, index, item)
            continue

        if block["type"] == "table":
            header = block["header"]
            rows = block["rows"]
            table = doc.add_table(rows=1, cols=len(header))
            for idx, value in enumerate(header):
                set_cell_text(table.rows[0].cells[idx], value, bold=True)
            for row in rows:
                row_cells = table.add_row().cells
                for idx, value in enumerate(row):
                    if idx < len(row_cells):
                        set_cell_text(row_cells[idx], value)
            format_table(table)
            doc.add_paragraph()


def build() -> Path:
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    cover_values, body_start = parse_cover_values(lines)
    blocks = parse_markdown(lines[body_start:])

    doc = Document(TEMPLATE_DOCX)
    fill_cover(doc, cover_values)
    clear_template_body(doc)
    render_markdown(doc, blocks)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build())
