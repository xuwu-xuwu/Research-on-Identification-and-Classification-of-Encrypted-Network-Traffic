from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORTS_DIR = ROOT / "docs" / "reports"
SOURCE_MD = REPORTS_DIR / "中期报告-网络流量加密方法识别与分类技术研究.md"
TEMPLATE_DOCX = next((ROOT / "docs" / "templates" / "信息系统安全与对抗技术-报告模板").glob("*302*中期报告*.docx"))
OUTPUT_DOCX = REPORTS_DIR / "中期报告-网络流量加密方法识别与分类技术研究-模板规范版.docx"

BODY_CN_FONT = "宋体"
BODY_EN_FONT = "Times New Roman"
BODY_FONT_SIZE = Pt(12)
SUBHEAD_CN_FONT = "黑体"
SUBHEAD_EN_FONT = "Times New Roman"
SUBHEAD_FONT_SIZE = Pt(14)
TABLE_FONT_SIZE = Pt(10.5)
FIRST_LINE_INDENT = Cm(0.74)


def strip_inline_markdown(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("**", "")
    text = text.replace("*", "")
    text = text.replace("\u00a0", " ")
    return text.strip()


def parse_pipe_row(line: str) -> list[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [strip_inline_markdown(part) for part in raw.split("|")]


def parse_cover_values(lines: list[str]) -> tuple[list[str], int]:
    values: list[str] = []
    if len(lines) < 3 or not lines[0].strip().startswith("|"):
        return values, 0
    index = 2
    while index < len(lines) and lines[index].strip().startswith("|"):
        row = parse_pipe_row(lines[index])
        if len(row) >= 2:
            values.append(row[1])
        index += 1
    return values, index


def starts_special(line: str) -> bool:
    stripped = line.strip()
    return bool(
        re.match(r"^(#{1,6})\s+", stripped)
        or stripped.startswith("|")
        or re.match(r"^\d+\.\s+", stripped)
    )


def parse_markdown(lines: list[str]) -> list[dict]:
    blocks: list[dict] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading_match.group(1)),
                    "text": strip_inline_markdown(heading_match.group(2)),
                }
            )
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            rows = [parse_pipe_row(row_line) for row_line in table_lines]
            if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in rows[1]):
                header = rows[0]
                body = rows[2:]
            else:
                header = rows[0]
                body = rows[1:]
            blocks.append({"type": "table", "header": header, "rows": body})
            continue

        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                items.append(strip_inline_markdown(item))
                i += 1
            blocks.append({"type": "olist", "items": items})
            continue

        para_lines = []
        while i < len(lines):
            current = lines[i].strip()
            if not current or starts_special(lines[i]):
                break
            para_lines.append(strip_inline_markdown(current))
            i += 1
        blocks.append({"type": "para", "text": " ".join(para_lines)})

    return blocks


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_run_font(run, cn_font: str, en_font: str, size: Pt, bold: bool | None = None) -> None:
    run.font.name = en_font
    run.font.size = size
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key, value in (
        ("w:ascii", en_font),
        ("w:hAnsi", en_font),
        ("w:eastAsia", cn_font),
        ("w:cs", en_font),
    ):
        rfonts.set(qn(key), value)


def set_cell_text(cell, text: str, align=WD_ALIGN_PARAGRAPH.CENTER, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, BODY_CN_FONT, BODY_EN_FONT, TABLE_FONT_SIZE, bold=bold)
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.left_indent = Cm(0)
    fmt.line_spacing = 1.2
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def fill_cover(doc: Document, values: list[str]) -> None:
    project_table = doc.tables[0]
    if len(project_table.rows) >= 1 and len(project_table.rows[0].cells) >= 2:
        set_cell_text(project_table.cell(0, 1), "302")

    info_table = doc.tables[1]
    for row_idx, value in enumerate(values):
        if row_idx < len(info_table.rows):
            set_cell_text(info_table.cell(row_idx, 1), value)


def update_note_paragraph(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("注："):
            paragraph.text = "注：中期报告（不少于3000字）"
            for run in paragraph.runs:
                set_run_font(run, BODY_CN_FONT, BODY_EN_FONT, BODY_FONT_SIZE)
            return


def clear_template_body(doc: Document) -> None:
    start_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "项目研究进展和任务完成情况":
            start_index = i
            break
    if start_index is None:
        raise ValueError("Could not locate the start of the template body.")

    for paragraph in list(doc.paragraphs[start_index:]):
        delete_paragraph(paragraph)


def format_body_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = FIRST_LINE_INDENT
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, BODY_CN_FONT, BODY_EN_FONT, BODY_FONT_SIZE)


def format_heading_paragraph(paragraph, level: int) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.left_indent = Cm(0)
    fmt.space_before = Pt(6 if level > 1 else 12)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.25
    font_size = SUBHEAD_FONT_SIZE if level > 1 else Pt(14)
    for run in paragraph.runs:
        set_run_font(run, SUBHEAD_CN_FONT, SUBHEAD_EN_FONT, font_size, bold=True)


def format_list_paragraph(paragraph, index: int, text: str) -> None:
    paragraph.text = ""
    run = paragraph.add_run(f"{index}. {text}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.left_indent = FIRST_LINE_INDENT
    fmt.first_line_indent = -FIRST_LINE_INDENT
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    set_run_font(run, BODY_CN_FONT, BODY_EN_FONT, BODY_FONT_SIZE)


def format_table(table) -> None:
    for style_name in ("Table Grid", "TableGrid", "网格型", "网络型"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fmt = paragraph.paragraph_format
                fmt.first_line_indent = Cm(0)
                fmt.left_indent = Cm(0)
                fmt.line_spacing = 1.2
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        BODY_CN_FONT,
                        BODY_EN_FONT,
                        TABLE_FONT_SIZE,
                        bold=(r_idx == 0),
                    )


def render_markdown(doc: Document, blocks: list[dict]) -> None:
    for block in blocks:
        if block["type"] == "heading":
            paragraph = doc.add_paragraph()
            paragraph.add_run(block["text"])
            format_heading_paragraph(paragraph, level=block["level"])
            continue

        if block["type"] == "para":
            paragraph = doc.add_paragraph()
            paragraph.add_run(block["text"])
            format_body_paragraph(paragraph)
            continue

        if block["type"] == "olist":
            for index, item in enumerate(block["items"], start=1):
                paragraph = doc.add_paragraph()
                format_list_paragraph(paragraph, index, item)
            continue

        if block["type"] == "table":
            header = block["header"]
            rows = block["rows"]
            table = doc.add_table(rows=1, cols=len(header))
            for idx, value in enumerate(header):
                set_cell_text(table.rows[0].cells[idx], value, bold=True)
            for row in rows:
                row_cells = table.add_row().cells
                for idx, value in enumerate(row):
                    if idx < len(row_cells):
                        set_cell_text(row_cells[idx], value, align=WD_ALIGN_PARAGRAPH.CENTER)
            format_table(table)


def build() -> Path:
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    cover_values, body_start = parse_cover_values(lines)
    blocks = parse_markdown(lines[body_start:])

    doc = Document(TEMPLATE_DOCX)
    fill_cover(doc, cover_values)
    update_note_paragraph(doc)
    clear_template_body(doc)
    render_markdown(doc, blocks)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build())
